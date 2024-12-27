import json
from flask import Flask, render_template, request, jsonify
from flask import Flask, request, redirect, url_for
import subprocess
import os
import sys
import logging
from flask import Flask, send_from_directory
import re
from upstash_redis import Redis
from docx import Document




file = open('keys.json', 'r')
keys = json.load(file)
upstash_token = keys['UPSTASH_REDIS_REST_TOKEN']
subscription_key = keys['subscription_key_1']
openai_api_key = keys['openai.api_key']

UPSTASH_REDIS_REST_URL="https://fine-swift-52766.upstash.io"
UPSTASH_REDIS_REST_TOKEN = upstash_token
redis_client = Redis(url=UPSTASH_REDIS_REST_URL, token=UPSTASH_REDIS_REST_TOKEN)

# Example file URLs (replace with your actual URLs)
mom_docx_file = "minutes_of_meeting.docx"
transcript_file= "transcript.docx"

app = Flask(__name__)
OUTPUT_DIR = r".\meeting assistant ai"
# Configure logging
logging.basicConfig(level=logging.DEBUG)

@app.route("/")
def home():
    """
    Render the home page with initial or processed meeting information.
    """
    meeting_info_json = request.args.get('meeting_info', default=None)

    if meeting_info_json:
        # Deserialize the JSON string into a dictionary
        meeting_info = json.loads(meeting_info_json)
    else:
        # Default empty meeting information
        meeting_info = {
            "date_time": " ",
            "attendees": 0,  # Default to 0 attendees
            "summary": " ",
            "transcript": " ",
            "follow_ups": " ",
            "action_items": " ",
        }

    return render_template("index.html", meeting_info=meeting_info)

def parse_minutes(minutes_text):
    """
    Parse the minutes of meeting text file into structured data.
    """
    lines = minutes_text.splitlines()
    meetinginfo = {
        "date_time": "",
        "summary": "",
        "follow_ups": "",
        "action_items": "",
    }
    for i, line in enumerate(lines):
        if line.startswith("Date Time:"):
            meetinginfo["date_time"] = line.split(":", 1)[1].strip()
        elif line.startswith("Short Summary"):
            summary_lines = []
            for j in range(i + 1, len(lines)):
                if lines[j].strip() == "":
                    break
                summary_lines.append(lines[j].strip())
            meetinginfo["summary"] = " ".join(summary_lines)
        elif line.startswith("Call to Action"):
            action_items = []
            for j in range(i + 1, len(lines)):
                if lines[j].strip() == "":
                    break
                action_items.append(lines[j].strip())
            meetinginfo["action_items"] = "\n".join(action_items)
        elif line.startswith("Follow-Ups"):
            follow_ups = []
            for j in range(i + 1, len(lines)):
                if lines[j].strip() == "":
                    break
                follow_ups.append(lines[j].strip())
            meetinginfo["follow_ups"] = "\n".join(follow_ups)

    return meetinginfo

num_of_attendees=redis_client.get('num_attendees')
diarized_text=redis_client.get('diarised-output')
minutes_text=redis_client.get('minutes_of_meeting')

def get_meeting_info_from_redis():
    """
    Retrieve the meeting information from Redis and return it as a dictionary.
    """
    num_of_attendees=redis_client.get('num_attendees')
    diarized_text=redis_client.get('diarised-output')
    minutes_text=redis_client.get('minutes_of_meeting')
    if minutes_text:
        parsed_minutes = parse_minutes(minutes_text)
    else:
        parsed_minutes = {}
    meeting_info = {
        "date_time": parsed_minutes.get("date_time") or "Unavailable",
        "attendees": num_of_attendees or "Unavailable",
        "summary": parsed_minutes.get('summary') or "Unavailable",
        "follow_ups": parsed_minutes.get('follow_ups') or "No follow up meetings scheduled",
        "action_items": parsed_minutes.get('action_items') or "Unavailable",
        "transcript": diarized_text or "Unavailable",
        "follow_ups_link": redis_client.get('meeting_link') or " "
    }
    redis_client.delete('meeting_link')
    return meeting_info

def save_transcript_to_docx(content):
    try:
        path = os.path.join(app.config['UPLOAD_FOLDER'], "transcript.docx")
        doc = Document()
        doc.add_heading("MeetMate", level=1)
        doc.add_heading("Meeting Transcript", level=2)

        doc.add_paragraph(content)
        doc.save(path)
        return path  # Return the file path for download
    except Exception as e:
        print(f"Error saving Meeting Transcript to DOCX: {e}")
        return None

def save_mom_to_docx(minutes_text):
    try:
        path = os.path.join(app.config['UPLOAD_FOLDER'], "minutes_of_meeting.docx")
        doc = Document()
        doc.add_heading("MeetMate", level=1)
        doc.add_heading("Meeting Minutes", level=2)

        # Process the MOM text into sections
        sections = {
            "Title": "",
            "Date Time": "",
            "No of Attendees": "",
            "Short Summary": "",
            "Discussed Points": [],
            "Detailed Speaker-wise Contribution & Conversation": "",
            "Call to Action": "",
            "Follow-Ups": ""
        }

        # Parse sections based on known headings
        current_heading = None
        for line in minutes_text.split("\n"):
            line = line.strip()
            if "Title:" in line:
                current_heading = "Title"
                sections[current_heading] = line.split("Title:")[1].strip()
            elif "Date Time:" in line:
                current_heading = "Date Time"
                sections[current_heading] = line.split("Date Time:")[1].strip()
            elif "No of Attendees:" in line:
                current_heading = "No of Attendees"
                sections[current_heading] = line.split("No of Attendees:")[1].strip()
            elif "Short Summary" in line:
                current_heading = "Short Summary"
            elif "Discussed Points" in line:
                current_heading = "Discussed Points"
            elif "Detailed Speaker-wise Contribution & Conversation" in line:
                current_heading = "Detailed Speaker-wise Contribution & Conversation"
            elif "Call to Action" in line:
                current_heading = "Call to Action"
            elif "Follow-Ups" in line:
                current_heading = "Follow-Ups"
            elif line:  # Only process non-empty lines
                if current_heading == "Discussed Points":
                    sections[current_heading].append(line)
                elif current_heading:
                    sections[current_heading] += " " + line

        # Add parsed content to the document
        doc.add_heading(f"Title: {sections['Title']}", level=3)
        doc.add_paragraph(f"Date Time: {sections['Date Time']}")
        doc.add_paragraph(f"No of Attendees: {sections['No of Attendees']}")

        doc.add_heading("Short Summary", level=3)
        doc.add_paragraph(sections["Short Summary"])

        doc.add_heading("Discussed Points", level=3)
        for point in sections["Discussed Points"]:
            doc.add_paragraph(point, style="Normal")

        doc.add_heading("Detailed Speaker-wise Contribution & Conversation", level=3)
        doc.add_paragraph(sections["Detailed Speaker-wise Contribution & Conversation"])

        doc.add_heading("Call to Action", level=3)
        doc.add_paragraph(sections["Call to Action"])

        doc.add_heading("Follow-Ups", level=3)
        doc.add_paragraph(sections["Follow-Ups"])
        doc.save(path)
        return path  # Return the file path for download

    except Exception as e:
        print(f"Error saving Minutes of Meeting to DOCX: {e}")
        return None

UPLOAD_FOLDER = '/tmp/uploads/'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/download_transcript')
def download_transcript():
    diarized_text = redis_client.get('diarised-output') or "Transcript not available"
    file_path = save_transcript_to_docx(diarized_text)
    if file_path:
        return send_from_directory(app.config['UPLOAD_FOLDER'], "transcript.docx", as_attachment=True)
    return "Failed to generate transcript file", 500

@app.route('/download_minutes')
def download_minutes():
    minutes_text = redis_client.get('minutes_of_meeting') or "Minutes of Meeting not available"
    file_path = save_mom_to_docx(minutes_text)
    if file_path:
        return send_from_directory(app.config['UPLOAD_FOLDER'], "minutes_of_meeting.docx", as_attachment=True)
    return "Failed to generate minutes of meeting file", 500



@app.route('/upload', methods=['POST'])
def upload_file():
    """
    Handle file upload for meeting audio files, supporting file paths and uploaded files.
    """
    # Check if a file was uploaded
    if "audio_file" in request.files:
        file = request.files["audio_file"]
        if file.filename == "":
            return jsonify({"error": "No selected file"}), 400

        if file:
            file_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
            file.save(file_path)
            meeting_info = process_audio_file(file_path)
            meeting_info_json = json.dumps(meeting_info)
            return redirect(url_for('home', meeting_info=meeting_info_json))
    
    # Check if a file path is provided
    file_path = request.form.get("file_path")
    if file_path and os.path.exists(file_path):
        meeting_info = process_audio_file(file_path)
        meeting_info_json = json.dumps(meeting_info)
        return redirect(url_for('home', meeting_info=meeting_info_json))

    return jsonify({"error": "No valid file or file path provided"}), 400

def process_audio_file(file_path):
    """
    Process the uploaded audio file and generate meeting details.
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file at {file_path} does not exist.")
        
        # Run the external script to process the file
        subprocess.run(["python", "diarize_MOM.py", file_path])
        meeting_info = get_meeting_info_from_redis()
        return meeting_info
    except Exception as e:
        logging.error(f"Error processing audio file: {e}")
        return {"error": f"Error processing audio file: {str(e)}"}

@app.route('/schedule', methods=['GET', 'POST'])
def schedule_meeting():
    if request.method == 'POST':
        # Process form data
        meeting_details = request.form['details']
        return f"Meeting scheduled: {meeting_details}"
    return render_template('schedule.html')

@app.route('/process', methods=['POST'])
def process():
    try:
        # Handle file upload
        uploaded_file = request.files.get('file')
        if uploaded_file:
            save_path = os.path.join(UPLOAD_FOLDER, uploaded_file.filename)
            uploaded_file.save(save_path)

            return jsonify({
                "message": "File uploaded and saved successfully.",
                "file_path": save_path,
                "status": "success"
            })

        # Handle other processing logic (e.g., running a script)
        current_dir = os.path.dirname(os.path.abspath(__file__))
        script_path = os.path.join(current_dir, "diarize_MOM.py")

        # Check if script exists
        if not os.path.exists(script_path):
            return jsonify({
                "message": "Script not found!",
                "error": f"Path '{script_path}' does not exist.",
                "status": "error"
            })

        # Run the test.py script
        logging.debug(f"Running script at {script_path}")
        result = subprocess.run(
            [sys.executable, script_path],
            capture_output=True,
            text=True
        )

        if result.returncode == 0:
            return jsonify({
                "message": "Script executed successfully!",
                "output": result.stdout.strip(),
                "status": "success"
            })
        else:
            return jsonify({
                "message": "Script execution failed!",
                "output": result.stderr.strip(),
                "status": "error"
            })
    except Exception as e:
        logging.error(f"Error: {str(e)}")
        return jsonify({
            "message": "An error occurred while processing.",
            "error": str(e),
            "status": "error"
        })

if __name__ == '__main__':
    app.run(debug=True)
